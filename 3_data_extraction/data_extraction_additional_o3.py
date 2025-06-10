import os
import re
import json
import copy
import glob
from openai import OpenAI
from dotenv import load_dotenv

import re
from urllib.parse import urljoin
import json
import os
import pandas as pd
import re
import base64
import numpy as np
import copy
import glob
from docx import Document

from openai import OpenAI
import json
import pandas as pd
import os
import glob
import numpy as np
from tqdm import tqdm

# Load environment variables from .env file
load_dotenv()

# Get API key from environment variables
api_key = os.getenv("OPENAI_API")
client = OpenAI(api_key=api_key)

def extract_sentence(pdf_name):
  # Construct the path to the JSON file associated with the given PDF
  json_name=pdf_name+"/structuredData.json"

  a = open(json_name)# Open and load the JSON file
  j = json.load(a)
  extracted_sentence=""# Initialize a string to store the extracted sentences
  previous_item_class=0# To track the class of the previous item and determine if a newline is needed

  # Iterate through each element in the JSON file
  for item in j["elements"] :
    # Check if the item is a heading
    if ("//Document/H" in item['Path']) or (re.search(r"//Document/Sect.*/H", item['Path'])):
      if previous_item_class!=1:
        extracted_sentence+="\n"
        previous_item_class=1
      if "Text" in item:
        if ("reference" == item["Text"].strip().lower()) or ("references" == item["Text"].strip().lower()) or ("appendix" == item["Text"].strip().lower()):
          extracted_sentence+="\n"
          break;
        extracted_sentence+=item["Text"]
        extracted_sentence+="\n"

    # Check if the item is a table and convert table data to text
    elif (re.search("//Document/Table.*" ,item['Path'])) or (re.search(r"//Document/Sect.*/Table", item['Path'])): # Replace table data with text
      if previous_item_class!=2:
        extracted_sentence+="\n"
        previous_item_class=2
      if "filePaths" in item:
        for t in item["filePaths"]:
          if ".xlsx" in t:
            extracted_sentence+=t.replace(".xlsx", "").split("/")[-1]
            extracted_sentence+="\n"
            t=pdf_name+"/"+t
            # Load the Excel file
            df_t = pd.read_excel(t)
            # Remove carriage returns from column names and data
            df_t.columns = [col.replace('_x000D_', '') for col in df_t.columns]
            df_t = df_t.replace('_x000D_', '', regex=True)
            # Get column names as a string
            column_names = ' '.join(df_t.columns)

            extracted_sentence+=df_t.to_csv(index=False)
            extracted_sentence+="\n"

    # Check if the item is a footnote
    elif (re.search("//Document/Footnote\.*",item['Path'])) or (re.search(r"//Document/Sect.*/Footnote", item['Path'])):
      if previous_item_class!=5:
        extracted_sentence+="\n"
        previous_item_class=5
      if "Text" in item:
        extracted_sentence+=item["Text"]
        extracted_sentence+="\n"

    # Check if the item is a paragraph
    elif ("//Document/P" in item['Path']) or (re.search(r"//Document/Sect.*/P", item['Path'])):
      if previous_item_class!=3:
        extracted_sentence+="\n"
        previous_item_class=3
      if not("/Figure" in item["Path"]):
        if "Text" in item:
          extracted_sentence+=item["Text"]

    # Check if the item is a list
    elif ("//Document/L" in item['Path']) or (re.search(r"//Document/Sect.*/L", item['Path'])):
      if previous_item_class!=4:
        extracted_sentence+="\n"
        previous_item_class=4
      if "/Lbl" in item["Path"]:
        extracted_sentence+="\n"
      if "Text" in item:
        extracted_sentence+=item["Text"]


  # Close the JSON file
  a.close()
  return extracted_sentence



def extract_sentence_text(pdf_name):
  # Construct the path to the JSON file associated with the given PDF
  json_name=pdf_name+"/structuredData.json"

  a = open(json_name)# Open and load the JSON file
  j = json.load(a)
  extracted_sentence=""# Initialize a string to store the extracted sentences
  previous_item_class=0 # To track the class of the previous item and determine if a newline is needed

  # Iterate through each element in the JSON file
  for item in j["elements"] :
    # Check if the item is a heading
    if ("//Document/H" in item['Path']) or (re.search(r"//Document/Sect.*/H", item['Path'])):
      if previous_item_class!=1:
        extracted_sentence+="\n"
        previous_item_class=1
      if "Text" in item:
        # Check for specific keywords to stop extracting sentences
        if ("reference" == item["Text"].strip().lower()) or ("references" == item["Text"].strip().lower()) or ("appendix" == item["Text"].strip().lower()):
          extracted_sentence+="\n"
          break;
        extracted_sentence+=item["Text"]
        extracted_sentence+="\n"

    # Check if the item is a paragraph
    elif ("//Document/P" in item['Path']) or (re.search(r"//Document/Sect.*/P", item['Path'])):
      if previous_item_class!=3:
        extracted_sentence+="\n"
        previous_item_class=3
      if not("/Figure" in item["Path"]):
        if "Text" in item:
          extracted_sentence+=item["Text"]

    # Check if the item is a list
    elif ("//Document/L" in item['Path']) or (re.search(r"//Document/Sect.*/L", item['Path'])):
      if previous_item_class!=4:
        extracted_sentence+="\n"
        previous_item_class=4
      if "/Lbl" in item["Path"]:
        extracted_sentence+="\n"
      if "Text" in item:
        extracted_sentence+=item["Text"]

    # Check if the item is a footnote
    elif (re.search("//Document/Footnote\.*",item['Path'])) or (re.search(r"//Document/Sect.*/Footnote", item['Path'])):
      if previous_item_class != 5:
        extracted_sentence+="\n"
        previous_item_class=5
      if "Text" in item:
        extracted_sentence+=item["Text"]
        extracted_sentence+="\n"

  # Close the JSON file
  a.close()
  return extracted_sentence


import base64
from mimetypes import guess_type

# Function to encode a local image into data URL
def local_image_to_data_url(image_path):
    # Guess the MIME type of the image based on the file extension
    mime_type, _ = guess_type(image_path)
    if mime_type is None:
        mime_type = 'application/octet-stream'  # Default MIME type if none is found

    # Read and encode the image file
    with open(image_path, "rb") as image_file:
        base64_encoded_data = base64.b64encode(image_file.read()).decode('utf-8')

    # Construct the data URL
    return f"data:{mime_type};base64,{base64_encoded_data}"

def GPT_chat_arm(tools, messages, tools_name):
  response = client.chat.completions.create(
        model="gpt-4o",
        messages=messages,
        tools=tools,
        tool_choice =  {"type": "function", "function": {"name": tools_name}},
        temperature=0
    )
  return response

def GPT_chat(tools, messages, tools_name):
  
    response = client.responses.create(
        model="o3", #use o3
        input=messages,
        text={ # This is for specifying the input text
            "format": {
                "type": "text"
            }
        },
        reasoning={ # This is for reasoning, 'high' has better precision (but consumes more tokens)
            "effort": "high"
        },
        tools=tools, # This is for tool specification, no change
        tool_choice={"type": "function", "name": tools_name},
        store=False, # If you have the same conversation multiple times, saving to storage reduces token count, but it doesn't matter if you're only asking once
        #temperature=0 # temperature cannot be used
    )
  
    return response


attention_query = """
#####Attention######

# From which paper to extract data
The primary paper is named AuthorYYYY, the registry ULR AuthorYYYY_NCT1234, the protocol AuthorYYYY_pr, the appendix/supplement files AuthorYYY_app or AuthorYYY_supp, and the secondary papers AuthorYYYY_s1,2,…
You can mainly extract data from primary paper.

#Mixed model repeated measures
- When to be careful
"mixed model", "mixed model repeated measures (MMRM)", "linear mixed model", "growth curve analysis ", "hierarchical model", etc.
If SEs are reported instead of SDs, be especially careful, as MMRM is highly likely.
- What to look out for in evaluation and data extraction
MMRM is said to be a method to estimate the effect of missing values of repeatedly measured outcomes without bias by making the assumption of missing at random
So, N (e.g., Severity_ep_n) in the extraction on outcomes should be the N that was analyzed (=in principle, number randomized; number analyzed for statistical analysis excluding early dropouts). Note that this is different from the number of people for whom outcome measures were available. (In the Christensen2016 example, 1149 patients were randomly assigned and 581 were available at 6 weeks for primary outcome measurement. (N should also be number randomized when calculating SD using SD=SE*sqrt(N).


# How to extract data for each arm
First, extract the data for each arm.
If the arm is further divided into several groups, merge these groups and calculate the integrated data.
If you cannot extract data by arm, extract data for all participants.
If you cannot find any data to extract or you do not have to extract data, output "-1".

# Which measure data you will take
You have to extract data (mean, SD, number, etc.) in the measure you previously extracted.

# The measure abbreviations
ISI: Insomnia Severity Index
SCI: Sleep Condition Indicator
FOSQ: Functional Outcomes of Sleep Questionnaire
ESS: Epworth Sleepiness Scale
PSQI: Pittsburgh Sleep Quality Index
AIS: Athens Insomnia Scale
WASO: Wake After Sleep Onset
SOL: Sleep Onset Latency

# How to calculate SD
You have to calculate SD and output as number (do not output as formula).
SD=SE*sqrt(n)
　If mixed model repeated measures were used, n is the number of people randomized or included in the analysis.
SD=(upper limit of [100-x]%CI - mean)*sqrt(n)/ NORM.S.INV([100-x/2]/100)
  If 90%CI, then ((upper limit of 90%CI - mean)*sqrt(n)/ norm.s.inv(0.95))
If 95%CI, then ((upper limit of 95%CI - mean)*sqrt(n)/ norm.s.inv(0.975))
If 99%CI, then ((upper limit of 99%CI - mean)*sqrt(n)/ norm.s.inv(0.995))
SD= (Q3 − Q1)/1.35
"""

eval_trials = {
    0: ["1", "1の関連論文"],
    1: ["3", "3の関連論文"],
    2: ["4"],
    3: ["5", "5_main_article", "5_protocol"],
    4: ["7", "7関連yoi160086", "7関連yoi160086supp2_prod"]
}


# Paper specification
shot_name = "chat"
log_dir_name = "extracted_data_log_GPT_o3"

print("OK")

n_index = 5
for fold_index in range(5,10):
  for eval_index in range(5):
    if (fold_index == 5) and (eval_index < 4):
      continue;
    eval_paper_list = eval_trials[eval_index]
    print("n_index: ", n_index)
    print("fold_index: ", fold_index)
    print("eval_index: ", eval_index)
    # For storing extracted data
    extracted_data = {}
    # For writing
    output_extracted_data = {}
############################################
    # Creating tools
    tools = [{'type': 'function',
        'function':
            {'name': 'extract_arm',
            'description': 'Extract the number of the arms and each arm name',
            'parameters': {
                'type': 'object',
              'properties': {
                  "arms_num":
                  {"type": "number",
                  "description": "Answer the number of the arms"
                  },
                  "arm_name_1":
                  {"type": "string",
                  "description": "Answer arm name. If exceed the number of the arms, answer 'nan'."
                  },

                  "arm_name_2":
                  {"type": "string",
                  "description": "Answer arm name. If exceed the number of the arms, answer 'nan'."
                  },

                  "arm_name_3":
                  {"type": "string",
                  "description": "Answer arm name. If exceed the number of the arms, answer 'nan'."
                  },

                  "arm_name_4":
                  {"type": "string",
                  "description": "Answer arm name. If exceed the number of the arms, answer 'nan'."
                  },

                  "arm_name_5":
                  {"type": "string",
                  "description": "Answer arm name. If exceed the number of the arms, answer 'nan'."
                  }
              },
            "required": [
                "arms_num",
                "arm_name_1",
                "arm_name_2",
                "arm_name_3",
                "arm_name_4",
                "arm_name_5"
                ]
            }
        }
    }
    ]
#######################################################
    # Creating messages for arm
    messages_arm=[]
    system_content = """You are a systematic reviewer.You have to extract arms(i.e., interventions and controls) from the inputted paper of a randomized controlled trial.\n
    Use the extract_arm tool in your response.\n\n"""
    user_content = """You have to extract arms(i.e., interventions and controls) from the inputted paper of a randomized controlled trial.\n
    Use the extract_arm tool in your response.\n\n"""
    png_files_sum = 0

    # Creating input text
    dirpath = f"fulltext/{eval_paper_list[0]}"
    print("OK1")
    main_paper_txt= extract_sentence_text(dirpath)
    user_content += main_paper_txt
    print("OK2")
    if main_paper_txt == "":
      # Generate log file name (including date and time)
      log_file_name = f"{log_dir_name}/{shot_name}/{n_index}_paper/{fold_index}_fold/error_log_{eval_index}_no_paper_text.txt"
      # Get error message
      error_message = f"An error occurred: no_paper_text"
      print(error_message)
      # Write error message to log file
      with open(log_file_name, 'w') as file:
          file.write(error_message)
      continue;

    messages_arm.append({"role": "system", "content": system_content})
    messages_arm.append({"role": "user", "content": user_content})

############################################
    # Getting response
    try:
      tools_name = "extract_arm"
      response_arm = GPT_chat_arm(tools, messages_arm, tools_name)
    except Exception as e:
      # Generate log file name (including date and time)
      log_file_name = f"{log_dir_name}/{shot_name}/{n_index}_paper/{fold_index}_fold/error_log_{eval_index}_response_arm.txt"
      # Get error message
      error_message = f"An error occurred: {str(e)}"
      print(error_message)
      # Write error message to log file
      with open(log_file_name, 'w') as file:
          file.write(error_message)
      continue;

############################################
    # Formatting response
    try:
      output_dict_arm = json.loads(response_arm.choices[0].message.tool_calls[0].function.arguments) # Output is in text format, split by line breaks
      arms_num = output_dict_arm["arms_num"]
      arms_list = []
      for arm_num in range(arms_num):
        arms_list.append(output_dict_arm[f"arm_name_{arm_num+1}"])
      print("arms_num: ", arms_num)
      print("arms_list: ", arms_list)
    except Exception as e:
      log_file_name = f"{log_dir_name}/{shot_name}/{n_index}_paper/{fold_index}_fold/error_log_{eval_index}_output_dict_arm.txt"
      # Get error message
      error_message = f"An error occurred: {str(e)}"
      print(error_message)
      # Write error message to log file
      with open(log_file_name, 'w') as file:
          file.write(error_message)
      continue;

############################################
    # Process each arm
    for arm_num in range(arms_num):
      print("arm_num: ", arm_num)
############################################
      # Create dictionary to store output
      # JSON file path
      json_file_path = 'original_description/original_description_insomnia.json'

      # Load JSON data from file
      with open(json_file_path, 'r') as file:
          data = json.load(file)

      # Extract numeric fields
      all_fields = [field for field in data]

      output = {item['name']: "None" for item in all_fields}

############################################
      # Create tools from optimized description
      log_path = f'optimized_description/{n_index}_paper/{fold_index}_fold/optimized_all_fields_{n_index-1}*.json'
      file_paths = glob.glob(log_path)

      for file_path in file_paths:
          with open(file_path, 'r') as file:
              optimized_all_fields = json.load(file)

      # For when options are missing
      for item in optimized_all_fields:
        if item['name'] == 'ind_clu':
          if not('[individual, cluster]' in item['description']):
            item['description'] += ' You have to choose from the options [individual, cluster].'
        if item['name'] == 'Single_multi':
          if not('[single, multi]' in item['description']):
            item['description'] += ' You have to choose from the options [single, multi].'
        if item['name'] == 'Insomnia_diagnosis':
          if not('[1. formal_DSM, 1. formal_ICSD, 1. formal_ICD, 1. formal_Edinger2004, 2. informal_ISI, 2. informal_SCI, 2. informal_PSQI, 2. informal_AIS, 2. informal_other_self_reported_scales, 3. others]' in item['description']):
            item['description'] += ' You have to choose from the options [1. formal_DSM, 1. formal_ICSD, 1. formal_ICD, 1. formal_Edinger2004, 2. informal_ISI, 2. informal_SCI, 2. informal_PSQI, 2. informal_AIS, 2. informal_other_self_reported_scales, 3. others].'
        if item['name'] == 'Comorbidities':
          if not('[0. primary_insomnia, 1. primary_secondary_mixed, 2. psy_MDD, 2. psy_anxiety, 2. psy_PTSD, 2. psy_schizophrenia_or_psychosis, 2. psy_bipolar, 2. psy_substance_dependence, 2. psy_others, 2. psy_mixed, 3. phy_dialysis, 3. phy_heart_failure, 3. phy_diabetes, 3. phy_fibromyalgia, 3. phy_cancer, 3. phy_steroid, 3. phy_pregnancy_perinatal, 3. phy_menopause, 3. phy_others, 3. phy_mixed, 4. psy_phy_mixed]' in item['description']):
            item['description'] += ' You have to choose from the options [0. primary_insomnia, 1. primary_secondary_mixed, 2. psy_MDD, 2. psy_anxiety, 2. psy_PTSD, 2. psy_schizophrenia_or_psychosis, 2. psy_bipolar, 2. psy_substance_dependence, 2. psy_others, 2. psy_mixed, 3. phy_dialysis, 3. phy_heart_failure, 3. phy_diabetes, 3. phy_fibromyalgia, 3. phy_cancer, 3. phy_steroid, 3. phy_pregnancy_perinatal, 3. phy_menopause, 3. phy_others, 3. phy_mixed, 4. psy_phy_mixed].'
        if item['name'] == 'scale_used':
          if not('[ISI, SCI, FOSQ, ESS, PSQI, AIS, other_self_reported_scales, WASO_SOL, WASO, SOL]' in item['description']):
            item['description'] += ' You have to choose from the options [ISI, SCI, FOSQ, ESS, PSQI, AIS, other_self_reported_scales, WASO_SOL, WASO, SOL].'
        if item['name'] == 'Severity_scale':
          if not('[ISI, SCI, FOSQ, ESS, PSQI, AIS, other_self_reported_scales, WASO_SOL, WASO, SOL]' in item['description']):
            item['description'] += ' You have to choose from the options [ISI, SCI, FOSQ, ESS, PSQI, AIS, other_self_reported_scales, WASO_SOL, WASO, SOL].'
        if item['name'] == 'r_long_scale_used':
          if not('[ISI, SCI, FOSQ, ESS, PSQI, AIS, other_self_reported_scales, WASO_SOL, WASO, SOL]' in item['description']):
            item['description'] += ' You have to choose from the options [ISI, SCI, FOSQ, ESS, PSQI, AIS, other_self_reported_scales, WASO_SOL, WASO, SOL].'

      properties_dict = {item['name']: {k: v for k, v in item.items() if k != 'name'} for item in optimized_all_fields}
      name_list = [item['name'] for item in optimized_all_fields]

      description_tools = [{'type': 'function',
              'name': 'extract_study_features',
              'description': 'Extracts key features from a user-inputted research article for systematic review purposes.',
              'parameters': {
                  'type': 'object',
                'properties': {}
              }
          }
      ]
      splitted_name_list = copy.deepcopy(name_list)
      description_tools[0]['parameters']['required'] = splitted_name_list
      description_tools[0]['parameters']['properties'] = {key: properties_dict[key] for key in splitted_name_list if key in properties_dict}

############################################
      # Creating messages
      system_content_DE = """You are a systematic reviewer.
      Use the extract_study_features tool in your response.
      """

      query_GPT = f"""You have to extract features in the {arms_list[arm_num]} arm from the inputted paper of a randomized controlled trial.
        Use the extract_study_features tool in your response.\n
          """
      query_GPT += attention_query

      query_function_calling = copy.deepcopy(query_GPT)

      query_GPT += "The extract_study_features:"
      query_GPT += json.dumps(description_tools)

      last_query = attention_query

      messages_GPT = []
      messages_function_calling = []
      content_prompt_GPT = []

      png_files_sum = 0

      content_prompt_GPT.append({"type": "input_text", "text": query_GPT})

      # Creating input text
      for paper_num in range(len(eval_paper_list)):
        dirpath = f"insomnia_DE_for_update/fulltext/{eval_paper_list[paper_num]}"
        paper_query = extract_sentence(dirpath)
        content_prompt_GPT.append({"type": "input_text", "text":paper_query})

        # List all files and directories in the directory, filter for PNG files only
        if os.path.exists(f"{dirpath}/figures"):
          png_files = [file for file in os.listdir(f"{dirpath}/figures") if file.endswith('.png')]

          for png_files_num in range(len(png_files)):
            if png_files_sum >= 10: # Limited to 10 images
              break;
            content_prompt_GPT.append({"type": "input_image", "image_url": local_image_to_data_url(f"{dirpath}/figures/"+png_files[png_files_num])})
            png_files_sum += 1

        csv_query = ""
        csv_flag = False
        # Process for reading CSV
        for file_name in os.listdir(dirpath):
            if file_name.endswith('.csv'):
                csv_flag = True
                csv_query += "docx_name: "
                csv_query += "docx_name: " + file_name + "\n"
                file_path = os.path.join(dirpath, file_name)
                # Read CSV file into dataframe
                data = pd.read_csv(file_path)
                # Convert to 'key: value \n' format
                for index, row in data.iterrows():
                    print(index)
                    for col in data.columns:
                        key = col.strip()
                        value = str(row[col]).strip()
                        csv_query += f"{key}: {value} \n"
        if csv_flag:
          content_prompt_GPT.append({"type": "input_text", "text":csv_query})


        docx_query = ""
        docx_flag = False
        # Process for reading txt, docx
        for file_name in os.listdir(dirpath):
            if file_name.endswith('.docx'):
                docx_flag = True
                docx_query += "docx_name: "
                docx_query += "docx_name: " + file_name + "\n"
                file_path = os.path.join(dirpath, file_name)
                doc = Document(file_path)
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)
                    # Read tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            full_text.append(cell.text)
                docx_query += '\n'.join(full_text)
        if docx_flag:
          content_prompt_GPT.append({"type": "input_text", "text":docx_query})


      content_prompt_GPT.append({"type": "input_text", "text": last_query})

      messages_GPT.append({"role": "system", "content": system_content_DE})
      messages_GPT.append({"role": "user", "content": content_prompt_GPT})

############################################
      try:
        # Getting response
        tools_name = "extract_study_features"
        response_GPT = GPT_chat(description_tools, messages_GPT ,tools_name)

############################################
        # Storing the results
        function_calls = [item for item in response_GPT.output if item.type == "function_call"]
        #output_dict = json.loads(response_GPT.output[0].arguments)
        output_dict = json.loads(function_calls[0].arguments)
        for k in output_dict.keys():
          if k in output.keys():
            output[k] = output_dict[k]
        extracted_data[arms_list[arm_num]] = copy.deepcopy(output)
      except Exception as e:
        print(response_GPT)
        extracted_data[arms_list[arm_num]] = copy.deepcopy(output)
        # Generate log file name (including date and time)
        log_file_name = f"{log_dir_name}/{shot_name}/{n_index}_paper/{fold_index}_fold/error_log_{eval_index}_data_extraction_armnum_{arm_num}.txt"
        # Get error message
        error_message = f"An error occurred: {str(e)}"
        print(error_message)
        # Write error message to log file
        with open(log_file_name, 'w') as file:
            file.write(error_message)
        continue;

############################################
    # Write extraction results of the evaluation paper to a JSON file
    file_name = f"{log_dir_name}/{shot_name}/{n_index}_paper/{fold_index}_fold/extracted_data_{eval_index}.json"
    output_extracted_data[eval_paper_list[0]] = copy.deepcopy(extracted_data)
    print("log has been created: ",extracted_data)
    with open(file_name, 'w') as file:
        json.dump(output_extracted_data, file, indent=4)
